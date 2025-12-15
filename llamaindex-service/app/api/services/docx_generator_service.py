from app.api.services.schema_generator_service import SchemaGeneratorService
from app.models.schema import (
    SchemaParagraph,
    SchemaNumberedList,
    SchemaBulletList,
    SchemaSubSubsection,
    SchemaSubsection,
    SchemaSection
)
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.text.paragraph import Paragraph
from typing import Union

class DocxGeneratorService:
    def __init__(self, schema_generator_service: SchemaGeneratorService):
        self.doc = Document()
        self.schema_generator_service = schema_generator_service
        
    def load_schema(self) -> None:
        for section in self.schema_generator_service.schema.sections:
            self.add_section(section=section)

    def add_section(self, section: SchemaSection) -> None:
        self.add_paragraph(section.heading)
        if section.description:
            self.add_paragraph(section.description)
        if len(section.subsections) > 0:
            for sub in section.subsections:
                self.add_subsection(subsection=sub)
        
    def add_subsection(self, subsection: SchemaSubsection) -> None:
        self.add_paragraph(subsection.heading)
        if subsection.description:
            self.add_paragraph(subsection.description)
        if len(subsection.subsubsections) > 0:
            for sub in subsection.subsubsections:
                self.add_subsubsection(subsubsection=sub)
        
    def add_subsubsection(self, subsubsection: SchemaSubSubsection) -> None:
        self.add_paragraph(subsubsection.heading)
        if subsubsection.description:
            self.add_paragraph(subsubsection.description)
        if len(subsubsection.elements) > 0:
            for element in subsubsection.elements:
                if isinstance(element, SchemaParagraph):
                    self.add_paragraph(paragraph=element)
                elif isinstance(element, SchemaNumberedList):
                    self.add_numbered_list(numbered_list=element)
                elif isinstance(element, SchemaBulletList):
                    self.add_bullet_list(bullet_list=element)
            
                
    def add_paragraph(self, paragraph: Union[SchemaParagraph, str]) -> None:
        if isinstance(paragraph, str):
            self.doc.add_paragraph(text=paragraph)
        else:
            self.doc.add_paragraph(text=paragraph.text)
    
    def add_numbered_list(self, numbered_list: SchemaNumberedList) -> None:
        if len(numbered_list.elements) == 0:
            raise ValueError("List `numbered_list` is empty.")
        self.doc.add_heading(numbered_list.heading.text, level=3)
        for idx, element in enumerate(numbered_list.elements):
            self.add_paragraph(f"{idx}. {element.text}")
            
    def add_bullet_list(self, bullet_list: SchemaBulletList) -> None:
        if len(bullet_list.elements) == 0:
            raise ValueError("List `bullet_list` is empty.")
        self.doc.add_heading(bullet_list.heading.text, level=3)
        for element in bullet_list.elements:
            self.add_paragraph(f"• {element.text}")
                    
    def insert_new_page(self) -> None:
        self.doc.add_page_break()
        
    def generate(self) -> None:
        # Zapis
        filename = '/app/generated_doc.docx'
        self.doc.save(filename)
        print(f"Wygenerowano kompletny szablon: {filename}")
        
    # Utils
    def remove_spacing(paragraph: Paragraph):
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0 
    