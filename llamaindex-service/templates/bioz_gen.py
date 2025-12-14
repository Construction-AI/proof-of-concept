from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_path_based_template():
    doc = Document()
    
    # Stylizacja dla czytelności
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Tytuł
    head = doc.add_heading('PLAN BIOZ', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- SEKCJA 1: DANE ---
    doc.add_heading('1. DANE INWESTYCJI', level=1)

    # Funkcja pomocnicza do dodawania pogrubionych etykiet
    def add_field(label, tag):
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        # Wstawiamy pełną ścieżkę w nawiasach klamrowych
        p.add_run(f'{{{{ {tag} }}}}') 

    # Używamy PEŁNYCH ścieżek z JSON-a
    add_field('Nazwa obiektu: ', 'content.title_page.object_info.name')
    add_field('Ulica: ', 'content.title_page.object_info.address.street')
    add_field('Miasto/Kod: ', 'content.title_page.object_info.address.zip_city')
    add_field('Działka: ', 'content.title_page.object_info.address.plot_number')
    
    doc.add_paragraph()
    add_field('Inwestor: ', 'content.title_page.investor.name')
    add_field('Adres Inwestora: ', 'content.title_page.investor.address')
    
    doc.add_paragraph()
    add_field('Kierownik Budowy: ', 'content.title_page.site_manager.name')
    add_field('Nr uprawnień: ', 'content.title_page.site_manager.qualifications_number')

    doc.add_page_break()

    # --- SEKCJA 2: OPIS ---
    doc.add_heading('2. CZĘŚĆ OPISOWA', level=1)
    
    doc.add_heading('2.1. Opis i Etapy', level=2)
    doc.add_paragraph('{{ content.descriptive_part.scope_of_works.description }}')

    doc.add_paragraph('Etapy robót:', style='List Bullet')
    
    # PĘTLA: Tutaj musimy iterować po liście ukrytej głęboko w strukturze
    # Zmienna iteracyjna to np. 'stage'
    doc.add_paragraph('{% for stage in content.descriptive_part.scope_of_works.stages %}')
    doc.add_paragraph('{{ stage }}', style='List Bullet')
    doc.add_paragraph('{% endfor %}')

    # --- SEKCJA 3: ZAGROŻENIA ---
    doc.add_heading('2.2. Zagrożenia', level=2)
    
    # Tabela zagrożeń
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Rodzaj zagrożenia'
    hdr[1].text = 'Kontekst'

    # Wiersz z pętlą (tr for) - używamy pełnej ścieżki do listy
    row = table.add_row().cells
    # Specyficzna składnia docxtpl dla tabel:
    row[0].text = '{% for h in content.descriptive_part.work_hazards %}{{ h.type }}'
    row[1].text = '{{ h.context }}{% endfor %}'

    # Zapis
    filename = 'szablon_bioz_paths.docx'
    doc.save(filename)
    print(f"Wygenerowano szablon: {filename}")

if __name__ == "__main__":
    create_path_based_template()