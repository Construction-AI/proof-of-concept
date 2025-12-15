from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

def create_full_bioz_template():
    doc = Document()

    # --- 1. FUNKCJE POMOCNICZE ---
    def remove_spacing(paragraph):
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0 

    def add_numbered_section(heading_text, jinja_path):
        p_if = doc.add_paragraph(f'{{% if {jinja_path} and {jinja_path} %}}')
        remove_spacing(p_if)
        doc.add_heading(heading_text, level=3)
        p_loop = doc.add_paragraph(f'{{% for item in {jinja_path} %}}')
        remove_spacing(p_loop)
        p_item = doc.add_paragraph('{{ item }}', style='List Number')
        p_item.paragraph_format.space_after = Pt(2) 
        p_end_loop = doc.add_paragraph('{% endfor %}')
        remove_spacing(p_end_loop)
        p_endif = doc.add_paragraph('{% endif %}')
        remove_spacing(p_endif)
        
    def add_bullet_section(heading_text, jinja_path):
        p_if = doc.add_paragraph(f'{{% if {jinja_path} and {jinja_path} %}}')
        remove_spacing(p_if)
        doc.add_heading(heading_text, level=3)
        p_loop = doc.add_paragraph(f'{{% for item in {jinja_path} %}}')
        remove_spacing(p_loop)
        p_item = doc.add_paragraph('{{ item }}', style='List Bullet')
        p_item.paragraph_format.space_after = Pt(2) 
        p_end_loop = doc.add_paragraph('{% endfor %}')
        remove_spacing(p_end_loop)
        p_endif = doc.add_paragraph('{% endif %}')
        remove_spacing(p_endif)

    # --- 2. KONFIGURACJA STYLÓW ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri Light'
        h_style.font.color.rgb = None
        h_style.font.bold = True
        if level == 1:
            h_style.font.size = Pt(16)
            h_style.paragraph_format.space_before = Pt(24)
            h_style.paragraph_format.space_after = Pt(12)

    # --- 3. STRONA TYTUŁOWA ---
    title = doc.add_heading('PLAN BEZPIECZEŃSTWA\nI OCHRONY ZDROWIA (BIOZ)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    table_cover = doc.add_table(rows=0, cols=2)
    table_cover.autofit = False
    table_cover.columns[0].width = Cm(5)
    table_cover.columns[1].width = Cm(11)

    def add_cover_row(label, tag_value):
        row = table_cover.add_row()
        remove_spacing(row.cells[0].paragraphs[0])
        remove_spacing(row.cells[1].paragraphs[0])
        
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        row.cells[1].paragraphs[0].add_run(f'{{{{ {tag_value} }}}}')
        return row

    add_cover_row('NAZWA OBIEKTU:', 'content.title_page.object_info.name')
    add_cover_row('ADRES INWESTYCJI:', 'content.title_page.object_info.address.zip_city')
    add_cover_row('INWESTOR:', 'content.title_page.investor.name')
    add_cover_row('ADRES INWESTORA:', 'content.title_page.investor.address')
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_cover_row('GENERALNY WYKONAWCA:', 'content.title_page.contractor.name')
    
    doc.add_paragraph().paragraph_format.space_after = Pt(40)

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig.add_run('Opracował:').font.size = Pt(10)
    
    doc.add_paragraph('_' * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- 4. TREŚĆ DOKUMENTU ---

    # 1. CZĘŚĆ OPISOWA
    doc.add_heading('1. CZĘŚĆ OPISOWA', level=1)
    
    doc.add_heading('1.1. Przedmiot inwestycji', level=2)
    doc.add_paragraph('{{ content.descriptive_part.scope_of_works.description }}')

    doc.add_heading('1.2. Uczestnicy procesu budowlanego', level=2)
    p = doc.add_paragraph()
    p.add_run('Kierownik Budowy: ').bold = True
    p.add_run('{{ content.title_page.site_manager.name }} ')
    p.add_run('{% if content.title_page.site_manager.qualifications_number %}'
              '(nr upr.: {{ content.title_page.site_manager.qualifications_number }})'
              '{% else %}(brak danych o uprawnieniach){% endif %}')

    # --- NOWA SEKCJA: PODSTAWA PRAWNA (WYPEŁNIACZ) ---
    doc.add_heading('1.3. Podstawa prawna opracowania', level=2)
    legal_text = (
        "Plan Bezpieczeństwa i Ochrony Zdrowia (BIOZ) został opracowany na podstawie:\n"
        "1. Ustawy z dnia 7 lipca 1994 r. – Prawo budowlane (t.j. Dz. U. z 2023 r. poz. 682 z późn. zm.).\n"
        "2. Rozporządzenia Ministra Infrastruktury z dnia 6 lutego 2003 r. w sprawie bezpieczeństwa i higieny pracy "
        "podczas wykonywania robót budowlanych (Dz. U. Nr 47, poz. 401).\n"
        "3. Dokumentacji projektowej oraz informacji dotyczącej bezpieczeństwa i ochrony zdrowia."
    )
    p_legal = doc.add_paragraph(legal_text)
    p_legal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    # 2. ZAKRES I KOLEJNOŚĆ ROBÓT
    doc.add_heading('2. ZAKRES I KOLEJNOŚĆ ROBÓT', level=1)
    doc.add_paragraph('Przewiduje się następującą kolejność realizacji robót budowlanych:')

    p_start = doc.add_paragraph("{% for stage in content.descriptive_part.scope_of_works.stages %}")
    remove_spacing(p_start)

    p_item = doc.add_paragraph("{{ loop.index }}. {{ stage }}")
    remove_spacing(p_item)

    p_end = doc.add_paragraph("{% endfor %}")
    remove_spacing(p_end)

    # 3. ZAGOSPODAROWANIE
    doc.add_heading('3. ZAGOSPODAROWANIE PLACU BUDOWY', level=1)
    
    doc.add_heading('3.1. Istniejące obiekty budowlane', level=2)
    add_bullet_section("Obiekty: ", "content.descriptive_part.existing_objects")

    doc.add_heading('3.2. Zagrożenia od infrastruktury technicznej', level=2)
    
    table_infra = doc.add_table(rows=1, cols=2)
    table_infra.style = 'Table Grid'
    hdr = table_infra.rows[0].cells
    hdr[0].text = 'Element infrastruktury'
    hdr[1].text = 'Opis / Zagrożenie / Zalecenia' # Rozszerzony nagłówek
    
    def add_infra_row(label, path):
        row = table_infra.add_row().cells
        remove_spacing(row[0].paragraphs[0])
        remove_spacing(row[1].paragraphs[0])
        row[0].text = label
        # Tu zakładamy, że model zwróci tekst zawsze, więc usuwam 'if/else Nie dotyczy'
        # Jeśli JSON będzie pusty, word wyświetli puste pole, co jest ok.
        row[1].text = f"{{{{ {path} }}}}"

    add_infra_row('Linie energetyczne', 'content.descriptive_part.site_hazards.power_lines')
    add_infra_row('Instalacje podziemne', 'content.descriptive_part.site_hazards.underground_utilities')
    add_infra_row('Ruch drogowy', 'content.descriptive_part.site_hazards.traffic')


    # 4. ZAGROŻENIA I PROFILAKTYKA
    doc.add_heading('4. IDENTYFIKACJA ZAGROŻEŃ I ŚRODKI PROFILAKTYCZNE', level=1)
    doc.add_paragraph('Na podstawie analizy dokumentacji projektowej oraz specyfiki robót zidentyfikowano następujące zagrożenia:')

    p_h_start = doc.add_paragraph('{% for hazard in content.descriptive_part.work_hazards %}')
    remove_spacing(p_h_start)

    p_h_item = doc.add_paragraph('{{ hazard }}', style='List Bullet')
    p_h_item.paragraph_format.space_after = Pt(2) 
    
    p_h_end = doc.add_paragraph('{% endfor %}')
    remove_spacing(p_h_end)

    # --- SZCZEGÓŁOWE INSTRUKCJE (BEZ IF-ów, ZAWSZE WIDOCZNE) ---
    doc.add_heading('4.1. Szczegółowe instrukcje bezpiecznego wykonywania robót', level=2)
    doc.add_paragraph('Poniżej przedstawiono szczegółowe wytyczne BHP dla kluczowych etapów robót:')

    add_numbered_section(
        'Zabezpieczenie wykopów i robót ziemnych', 
        'content.descriptive_part.preventive_measures.excavations'
    )

    add_numbered_section(
        'Praca na wysokości', 
        'content.descriptive_part.preventive_measures.heights'
    )

    add_numbered_section(
        'Praca maszyn budowlanych i transport', 
        'content.descriptive_part.preventive_measures.machinery'
    )

    # --- 5. STRONA Z DEBUG INFORMACJAMI ---
    doc.add_page_break()
    doc.add_heading('INFORMACJE DEBUG / GENERACJI', level=1)
    debug_table = doc.add_table(rows=0, cols=2)
    debug_table.autofit = False
    debug_table.columns[0].width = Cm(6)
    debug_table.columns[1].width = Cm(10)

    def add_debug_row(label, value):
        row = debug_table.add_row()
        remove_spacing(row.cells[0].paragraphs[0])
        remove_spacing(row.cells[1].paragraphs[0])
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)
        
    add_debug_row("Typ dokumentu: ", "{{ meta.document_type }}")
    add_debug_row("ID Firmy: ", "{{ meta.company_id }}")
    add_debug_row("ID Projektu: ", "{{ meta.project_id }}")
    add_debug_row("Autor: ", "{{ meta.author }}")
    add_debug_row("System Prompt: ", "{{ meta.system_instruction }}.")
    add_debug_row("Wersja: ", "{{ meta.version }}")
    add_debug_row("Język: ", "{{ meta.language }}")
    add_debug_row("Data utworzenia", "{{ meta.date_created }}")

    # Zapis
    filename = 'health_and_safety_plan.docx'
    doc.save(filename)
    print(f"Wygenerowano kompletny szablon: {filename}")

if __name__ == "__main__":
    create_full_bioz_template()