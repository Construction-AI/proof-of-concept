# app/services/document_export.py

import os
import logging
from typing import Dict, Any, Optional
from pathlib import Path
import datetime

# For DOCX generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# For XLSX generation
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter

logger = logging.getLogger(__name__)

class DocumentExportService:
    def __init__(self, export_dir: str = "/data/exports"):
        """
        Initialize document export service
        
        Args:
            export_dir: Directory for exported documents
        """
        self.export_dir = export_dir
        os.makedirs(self.export_dir, exist_ok=True)
    
    def export_document(self, document_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Export document to the appropriate format
        
        Args:
            document_data: Document data from the generator
            
        Returns:
            Dictionary with export result
        """
        try:
            document_format = document_data.get("format", "docx").lower()
            project_name = document_data.get("project_name", "project")
            document_type = document_data.get("document_type", "document")
            
            # Generate filename with timestamp
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            sanitized_name = self._sanitize_filename(project_name)
            filename = f"{timestamp}_{sanitized_name}_{self._sanitize_filename(document_type)}"
            
            if document_format == "docx":
                filepath = self._export_to_docx(document_data, f"{filename}.docx")
            elif document_format == "xlsx":
                filepath = self._export_to_xlsx(document_data, f"{filename}.xlsx")
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported document format: {document_format}"
                }
            
            return {
                "status": "completed",
                "filepath": filepath,
                "filename": os.path.basename(filepath),
                "message": f"Document exported successfully as {document_format.upper()}"
            }
            
        except Exception as e:
            logger.error(f"Error exporting document: {str(e)}")
            return {
                "status": "error",
                "message": f"Error exporting document: {str(e)}"
            }
    
    def _sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename to avoid invalid characters
        
        Args:
            filename: Original filename
            
        Returns:
            Sanitized filename
        """
        # Replace invalid characters with underscore
        invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        return filename
    
    def _export_to_docx(self, document_data: Dict[str, Any], filename: str) -> str:
        """
        Export document to DOCX format
        
        Args:
            document_data: Document data
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        doc = Document()
        
        # Add document title
        title = doc.add_heading(document_data.get("project_name", "Project"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document type as subtitle
        subtitle = doc.add_heading(document_data.get("document_type", "Document"), 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f"Data: {datetime.datetime.now().strftime('%d.%m.%Y')}")
        
        doc.add_paragraph()  # Add space
        
        # Add sections
        sections = document_data.get("sections", {})
        for section_key, section_data in sections.items():
            # Add section heading
            doc.add_heading(section_data.get("title", "Section"), 2)
            
            # Add section content
            content = section_data.get("content", "")
            if content:
                paragraph = doc.add_paragraph(content)
            else:
                paragraph = doc.add_paragraph("Brak danych")
                paragraph.italic = True
            
            doc.add_paragraph()  # Add space between sections
        
        # Save the document
        filepath = os.path.join(self.export_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _export_to_xlsx(self, document_data: Dict[str, Any], filename: str) -> str:
        """
        Export document to XLSX format
        
        Args:
            document_data: Document data
            filename: Output filename
            
        Returns:
            Path to exported file
        """
        wb = openpyxl.Workbook()
        ws = wb.active
        
        # Set title
        ws.title = "Document"
        
        # Add headers
        ws['A1'] = document_data.get("project_name", "Project")
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:D1')
        
        ws['A2'] = document_data.get("document_type", "Document")
        ws['A2'].font = Font(size=14, bold=True)
        ws.merge_cells('A2:D2')
        
        ws['A3'] = f"Data: {datetime.datetime.now().strftime('%d.%m.%Y')}"
        ws['A3'].alignment = Alignment(horizontal='right')
        ws.merge_cells('A3:D3')
        
        # Style for headers
        header_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
        
        # Add sections
        row = 5
        sections = document_data.get("sections", {})
        
        # Create section headers based on document type
        document_type = document_data.get("document_type", "")
        
        if "Karty materiałowe" in document_type:
            # Special format for material cards
            ws['A5'] = "Nazwa materiału"
            ws['B5'] = "Producent/dostawca"
            ws['C5'] = "Parametry techniczne"
            ws['D5'] = "Certyfikaty"
            ws['E5'] = "Warunki stosowania"
            ws['F5'] = "Metody montażu"
            
            for col in ['A', 'B', 'C', 'D', 'E', 'F']:
                ws[f'{col}5'].font = Font(bold=True)
                ws[f'{col}5'].fill = header_fill
            
            row = 6
            
            # Extract materials from sections
            for section_key, section_data in sections.items():
                content = section_data.get("content", "")
                if content and section_key == "material_name":
                    # Process material content (simplified)
                    materials = content.split("\n")
                    for material in materials:
                        material = material.strip()
                        if material:
                            ws[f'A{row}'] = material
                            row += 1
        
        elif "Zestawienia materiałowe" in document_type:
            # Special format for material lists and schedules
            ws['A5'] = "Materiał"
            ws['B5'] = "Jednostka"
            ws['C5'] = "Ilość"
            ws['D5'] = "Uwagi"
            
            for col in ['A', 'B', 'C', 'D']:
                ws[f'{col}5'].font = Font(bold=True)
                ws[f'{col}5'].fill = header_fill
            
            row = 6
            
            # Add a second sheet for schedule
            ws_schedule = wb.create_sheet(title="Harmonogram")
            ws_schedule['A1'] = "Harmonogram robót"
            ws_schedule['A1'].font = Font(size=14, bold=True)
            ws_schedule.merge_cells('A1:E1')
            
            ws_schedule['A3'] = "Etap"
            ws_schedule['B3'] = "Zadanie"
            ws_schedule['C3'] = "Rozpoczęcie"
            ws_schedule['D3'] = "Zakończenie"
            ws_schedule['E3'] = "Status"
            
            for col in ['A', 'B', 'C', 'D', 'E']:
                ws_schedule[f'{col}3'].font = Font(bold=True)
                ws_schedule[f'{col}3'].fill = header_fill
            
        else:
            # Default format - sections as headers
            ws['A5'] = "Sekcja"
            ws['B5'] = "Treść"
            
            ws['A5'].font = Font(bold=True)
            ws['B5'].font = Font(bold=True)
            ws['A5'].fill = header_fill
            ws['B5'].fill = header_fill
            
            row = 6
            
            # Add each section
            for section_key, section_data in sections.items():
                ws[f'A{row}'] = section_data.get("title", "Section")
                ws[f'B{row}'] = section_data.get("content", "")
                
                # Format cells
                ws[f'A{row}'].font = Font(bold=True)
                ws[f'B{row}'].alignment = Alignment(wrap_text=True)
                
                # Auto-adjust row height based on content
                ws.row_dimensions[row].height = max(15, len(section_data.get("content", "").split("\n")) * 15)
                
                row += 1
        
        # Adjust column widths
        for col in ['A', 'B', 'C', 'D', 'E', 'F']:
            if col in ws.column_dimensions:
                ws.column_dimensions[col].width = 20
        
        # Save the workbook
        filepath = os.path.join(self.export_dir, filename)
        wb.save(filepath)
        
        return filepath